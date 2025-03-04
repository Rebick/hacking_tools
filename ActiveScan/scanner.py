import subprocess
import re
import argparse
from docx import Document

def run_nmap_scan(target, verbose=False):
    """
    Ejecuta un escaneo de Nmap en dos fases y devuelve la salida completa.
    """
    command = f"""
    ports=$(nmap -p- --min-rate=1000 -Pn -T4 {target} | grep '^[0-9]' | cut -d '/' -f 1 | tr '\n' ',' | sed s/,$//) && \
    nmap -p$ports -Pn -sC -sV {target}
    """
    if verbose:
        print(f"Ejecutando: {command}")
    result = subprocess.run(command, shell=True, capture_output=True, text=True)
    return result.stdout

def parse_nmap_output(output, verbose=False):
    """
    Extrae la información relevante del escaneo de Nmap.
    """
    report = {}
    
    # Extraer dirección IP y hostname
    ip_match = re.search(r'Nmap scan report for (.*?) \((.*?)\)', output)
    if ip_match:
        report['hostname'] = ip_match.group(1)
        report['ip'] = ip_match.group(2)
    
    # Extraer dominios adicionales
    domains = re.findall(r'DNS:(\S+)', output)
    report['domains'] = domains if domains else []
    
    # Extraer puertos abiertos y servicios
    ports = re.findall(r'(\d+/tcp)\s+open\s+(\S+)\s+(.*)', output)
    report['ports'] = [
        {'port': port.split('/')[0], 'service': service, 'version': version}
        for port, service, version in ports
    ]
    
    # Intentar determinar el sistema operativo
    os_match = re.search(r'Service Info: Host: (.*?); OS: (.*?);', output)
    if os_match:
        report['os'] = os_match.group(2)
    else:
        report['os'] = "Desconocido"
    
    if verbose:
        print(f"Sistema Operativo detectado: {report['os']}")
        print(f"Cantidad de puertos abiertos: {len(report['ports'])}")
        if 'hostname' in report:
            print(f"Dominio principal encontrado: {report['hostname']}")
        if report['domains']:
            print(f"Dominios adicionales detectados: {', '.join(report['domains'])}")
    
    return report

def generate_word_report(report):
    """
    Genera un reporte en formato Word.
    """
    doc = Document()
    doc.add_heading('Reporte de Escaneo Nmap', 0)
    
    doc.add_heading('Detalles', level=1)
    doc.add_paragraph(f"Hostname: {report.get('hostname', 'Desconocido')}")
    doc.add_paragraph(f"Dirección IP: {report.get('ip', 'Desconocida')}")
    doc.add_paragraph(f"Sistema Operativo: {report.get('os', 'Desconocido')}")
    if report['domains']:
        doc.add_paragraph(f"Dominios adicionales detectados: {', '.join(report['domains'])}")
    
    doc.add_heading('Puertos Encontrados', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Puerto'
    hdr_cells[1].text = 'Servicio'
    hdr_cells[2].text = 'Versión'
    hdr_cells[3].text = 'Vector de Ataque'
    
    attack_vectors = {
        'smb': 'Explotación con SMBGhost, EternalBlue',
        'mysql': 'Credenciales por defecto, inyección SQL',
        'ldap': 'Extracción de información con ldapsearch',
        'kerberos': 'Ataques Pass-the-Ticket, Kerberoasting',
        'mssql': 'Ataques con Metasploit, credenciales débiles',
        'rpc': 'Enumeración con rpcclient',
        'http': 'Pruebas con Nikto, Burp Suite',
        'microsoft-ds': 'Explotación con SMBGhost, EternalBlue'  # Puerto 445 agregado
    }
    
    for port_info in report['ports']:
        row_cells = table.add_row().cells
        row_cells[0].text = port_info['port']
        row_cells[1].text = port_info['service']
        row_cells[2].text = port_info['version']
        row_cells[3].text = attack_vectors.get(port_info['service'].lower(), 'N/A')
    
    doc.save('nmap_report.docx')

def main():
    parser = argparse.ArgumentParser(description="Escaneo de red con Nmap")
    parser.add_argument("target", help="Dirección IP o dominio a escanear")
    parser.add_argument("-v", "--verbose", action="store_true", help="Modo detallado de salida")
    args = parser.parse_args()
    
    print("Iniciando escaneo...")
    scan_output = run_nmap_scan(args.target, args.verbose)
    print("Escaneo completado. Analizando datos...")
    report = parse_nmap_output(scan_output, args.verbose)
    
    print("Generando reporte...")
    generate_word_report(report)
    print("Reporte guardado como nmap_report.docx")

if __name__ == "__main__":
    main()
